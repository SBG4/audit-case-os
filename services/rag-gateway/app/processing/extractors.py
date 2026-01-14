"""
Text extraction from various document formats.

Supports:
- PDF files (PyPDF2)
- DOCX files (python-docx)
- Plain text files
- HTML files (BeautifulSoup)
- Future: Images with OCR (pytesseract)
"""

import logging
import mimetypes
from typing import Optional
from io import BytesIO

import PyPDF2
from docx import Document
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


class TextExtractor:
    """Base class for text extraction from different file formats."""

    @staticmethod
    def detect_mime_type(filename: str, content: bytes) -> str:
        """
        Detect MIME type from filename and content.

        Args:
            filename: Original filename
            content: File content as bytes

        Returns:
            MIME type string (e.g., "application/pdf")
        """
        # Try to guess from filename first
        mime_type, _ = mimetypes.guess_type(filename)

        if mime_type:
            return mime_type

        # Fallback to magic bytes detection
        if content.startswith(b'%PDF'):
            return "application/pdf"
        elif content.startswith(b'PK\x03\x04'):  # ZIP-based formats (DOCX, etc.)
            if b'word/' in content[:1000]:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return "application/zip"
        elif content.startswith(b'<'):  # Likely HTML/XML
            return "text/html"

        # Default to plain text
        return "text/plain"

    @staticmethod
    def extract_from_pdf(content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            content: PDF file content as bytes

        Returns:
            Extracted text

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")
                    continue

            if not text_parts:
                raise TextExtractionError("No text extracted from PDF")

            return "\n\n".join(text_parts)

        except PyPDF2.PdfReadError as e:
            raise TextExtractionError(f"Invalid PDF file: {e}")
        except Exception as e:
            raise TextExtractionError(f"PDF extraction failed: {e}")

    @staticmethod
    def extract_from_docx(content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            content: DOCX file content as bytes

        Returns:
            Extracted text

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            docx_file = BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            if not text_parts:
                raise TextExtractionError("No text extracted from DOCX")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise TextExtractionError(f"DOCX extraction failed: {e}")

    @staticmethod
    def extract_from_txt(content: bytes) -> str:
        """
        Extract text from plain text file.

        Args:
            content: Text file content as bytes

        Returns:
            Decoded text

        Raises:
            TextExtractionError: If decoding fails
        """
        # Try common encodings
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        raise TextExtractionError("Failed to decode text file with any supported encoding")

    @staticmethod
    def extract_from_html(content: bytes) -> str:
        """
        Extract text from HTML file.

        Args:
            content: HTML file content as bytes

        Returns:
            Extracted text with HTML tags removed

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            # Try to decode
            html_text = TextExtractor.extract_from_txt(content)

            # Parse HTML and extract text
            soup = BeautifulSoup(html_text, 'lxml')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            if not text.strip():
                raise TextExtractionError("No text extracted from HTML")

            return text

        except Exception as e:
            raise TextExtractionError(f"HTML extraction failed: {e}")


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract text from a file based on its MIME type.

    Args:
        filename: Original filename
        content: File content as bytes

    Returns:
        Extracted text

    Raises:
        TextExtractionError: If extraction fails or format is unsupported
    """
    mime_type = TextExtractor.detect_mime_type(filename, content)

    logger.info(f"Extracting text from {filename} (MIME: {mime_type})")

    try:
        if mime_type == "application/pdf":
            return TextExtractor.extract_from_pdf(content)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return TextExtractor.extract_from_docx(content)
        elif mime_type == "text/html":
            return TextExtractor.extract_from_html(content)
        elif mime_type.startswith("text/"):
            return TextExtractor.extract_from_txt(content)
        else:
            raise TextExtractionError(f"Unsupported file type: {mime_type}")

    except TextExtractionError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {filename}: {e}")
        raise TextExtractionError(f"Text extraction failed: {e}")
